"""翻译、导出、思维导图、搜索、代码生成、对比、课题组讨论等功能路由。"""
import os
import re
import sys
import json
import asyncio
import tempfile
import datetime
import logging

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse, FileResponse

from app.state import app_state
from app.dependencies import (
    UPLOAD_DIR,
    get_coordinator,
    get_lab_session_cache,
    set_lab_session_cache,
)
from app.schemas import (
    TranslateTextRequest,
    ExportRequest,
    SearchRequest,
    CompareRequest,
    LabDiscussRequest,
)
from agents.lab.lab_session import LabSession

logger = logging.getLogger(__name__)

router = APIRouter()


# ===================== 划词翻译 =====================

@router.post("/api/translate/text")
def translate_text(request: TranslateTextRequest):
    """划词翻译 - 翻译选中的文本片段。"""
    text = request.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="文本不能为空")
    if len(text) > 500:
        text = text[:500]

    coordinator = get_coordinator()
    try:
        result = coordinator.llm_service.chat_sync(
            user_message=f"请将以下英文翻译成中文，只输出翻译结果，不要解释：\n\n{text}",
            system_prompt="你是专业的学术论文翻译员。翻译要准确、专业、通顺。对于专业术语，在翻译后用括号附上英文原文。",
            chat_history=[],
        )
        return {"translation": result.strip(), "original": text}
    except Exception as e:
        logger.error(f"划词翻译失败: {e}")
        raise HTTPException(status_code=500, detail=f"翻译失败: {str(e)}")


# ===================== 导出报告 =====================

@router.post("/api/export/report")
def export_report(request: ExportRequest = ExportRequest()):
    """导出分析报告为 Word 文档。"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not app_state.current_summary:
        raise HTTPException(status_code=400, detail="暂无分析结果可导出")

    doc_info = app_state.document_info or {}
    summary = app_state.current_summary

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    # ---- 封面 ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 30)
    run.font.color.rgb = RGBColor(14, 165, 233)
    run.font.size = Pt(14)

    title_text = doc_info.get("title") or doc_info.get("filename", "论文分析报告")
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI 智能分析报告")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 116, 139)
    run.font.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.space_after = Pt(6)
    parts = []
    if doc_info.get("page_count"):
        parts.append(f"📄 {doc_info['page_count']} 页")
    if doc_info.get("word_count"):
        parts.append(f"📝 {doc_info['word_count']:,} 字")
    parts.append(f"🕐 {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run = meta.add_run("  |  ".join(parts))
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(148, 163, 184)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 30)
    run.font.color.rgb = RGBColor(14, 165, 233)
    run.font.size = Pt(14)

    doc.add_page_break()

    # ---- 目录 ----
    toc_title = doc.add_heading("目录", level=1)
    for run in toc_title.runs:
        run.font.color.rgb = RGBColor(14, 165, 233)

    h2_titles = re.findall(r'^##\s+(.+)$', summary, re.MULTILINE)
    if h2_titles:
        for i, t in enumerate(h2_titles):
            clean = re.sub(r'[#*`]', '', t).strip()
            p = doc.add_paragraph(f"{i+1}. {clean}", style='List Number')
            p.paragraph_format.space_after = Pt(4)
    else:
        doc.add_paragraph("（AI 自动生成分析内容）")

    doc.add_page_break()

    # ---- 正文 ----
    section_title = doc.add_heading("AI 分析报告", level=1)
    for run in section_title.runs:
        run.font.color.rgb = RGBColor(14, 165, 233)

    _render_markdown_to_docx(doc, summary)

    # ---- 标注 ----
    if request.annotations:
        doc.add_page_break()
        anno_title = doc.add_heading("📝 阅读标注", level=1)
        for run in anno_title.runs:
            run.font.color.rgb = RGBColor(251, 191, 36)

        doc.add_paragraph(f"共 {len(request.annotations)} 条标注")

        for i, anno in enumerate(request.annotations):
            doc.add_heading(f"标注 {i+1}", level=3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f"「{anno.get('text', '')}」")
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 116, 139)
            if anno.get("note"):
                note_p = doc.add_paragraph()
                note_p.paragraph_format.left_indent = Cm(1)
                run = note_p.add_run(f"📌 笔记: {anno['note']}")
                run.font.size = Pt(10)
            if anno.get("timestamp"):
                try:
                    ts = datetime.datetime.fromisoformat(anno["timestamp"].replace("Z", "+00:00"))
                    time_p = doc.add_paragraph()
                    time_p.paragraph_format.left_indent = Cm(1)
                    run = time_p.add_run(ts.strftime("%Y-%m-%d %H:%M"))
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(148, 163, 184)
                except Exception:
                    pass

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.space_before = Pt(24)
    run = footer_p.add_run("— 由 PaperReader 智能论文阅读助手生成 —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(148, 163, 184)
    run.font.italic = True

    filename = (doc_info.get("filename", "report") or "report").replace(".pdf", "").replace(".docx", "")
    output_path = os.path.join(tempfile.gettempdir(), f"{filename}_analysis_report.docx")
    doc.save(output_path)

    return FileResponse(
        path=output_path,
        filename=f"{filename}_analysis_report.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _render_markdown_to_docx(doc, md_text: str):
    """将 Markdown 文本渲染到 Word 文档。"""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = md_text.split("\n")
    in_code_block = False
    code_lines = []

    for line in lines:
        if line.strip().startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 116, 139)
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("#### "):
            h = doc.add_heading(re.sub(r'[#]', '', stripped[5:]).strip(), level=4)
            h.paragraph_format.space_before = Pt(8)
        elif stripped.startswith("### "):
            h = doc.add_heading(re.sub(r'[#]', '', stripped[4:]).strip(), level=3)
            h.paragraph_format.space_before = Pt(10)
        elif stripped.startswith("## "):
            h = doc.add_heading(re.sub(r'[#]', '', stripped[3:]).strip(), level=2)
            h.paragraph_format.space_before = Pt(14)
        elif stripped.startswith("# "):
            h = doc.add_heading(re.sub(r'[#]', '', stripped[2:]).strip(), level=1)
            h.paragraph_format.space_before = Pt(16)
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(stripped[2:])
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 116, 139)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, stripped[2:])
        elif re.match(r'^\d+\.\s', stripped):
            content = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style="List Number")
            _add_rich_text(p, content)
        elif stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 40)
            run.font.color.rgb = RGBColor(203, 213, 225)
            run.font.size = Pt(8)
        else:
            p = doc.add_paragraph()
            _add_rich_text(p, stripped)


def _add_rich_text(paragraph, text: str):
    """解析 Markdown inline 格式（粗体、斜体、行内代码）。"""
    from docx.shared import Pt, RGBColor

    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(14, 165, 233)
        else:
            paragraph.add_run(part)


# ===================== 智能搜索 =====================

@router.post("/api/search")
async def search_papers(request: SearchRequest):
    """搜索相关论文（Semantic Scholar API）。"""
    import httpx

    query = request.query.strip()

    if not query:
        if app_state.document_info:
            query = app_state.document_info.get("title", "") or app_state.document_info.get("filename", "")
        if not query:
            raise HTTPException(status_code=400, detail="请提供搜索关键词或先上传文档")

    try:
        data = None
        async with httpx.AsyncClient(timeout=15.0) as client:
            for attempt in range(3):
                resp = await client.get(
                    "https://api.semanticscholar.org/graph/v1/paper/search",
                    params={
                        "query": query[:200],
                        "limit": min(request.limit, 20),
                        "fields": "title,abstract,authors,year,citationCount,url,openAccessPdf,externalIds",
                    },
                )
                if resp.status_code == 429:
                    wait = 2 ** attempt
                    logger.warning(f"Semantic Scholar 429, 等待 {wait}s 重试...")
                    await asyncio.sleep(wait)
                    continue
                resp.raise_for_status()
                data = resp.json()
                break

        if data is None:
            raise HTTPException(status_code=429, detail="搜索请求过于频繁，请稍后再试")

        papers = []
        for p in data.get("data", []):
            authors = [a.get("name", "") for a in (p.get("authors") or [])[:5]]
            pdf_url = ""
            if p.get("openAccessPdf"):
                pdf_url = p["openAccessPdf"].get("url", "")
            doi = ""
            if p.get("externalIds"):
                doi = p["externalIds"].get("DOI", "")

            papers.append({
                "title": p.get("title", ""),
                "authors": authors,
                "year": p.get("year"),
                "citationCount": p.get("citationCount", 0),
                "abstract": (p.get("abstract") or "")[:300],
                "url": p.get("url", ""),
                "pdfUrl": pdf_url,
                "doi": doi,
            })

        return {"query": query, "total": data.get("total", 0), "papers": papers}

    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="搜索超时，请稍后重试")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"搜索失败: {e}")
        raise HTTPException(status_code=500, detail=f"搜索失败: {str(e)}")


# ===================== 全文翻译（pdf2zh）=====================

@router.post("/api/translate/stream")
async def translate_stream():
    """使用 pdf2zh 翻译论文 PDF（SSE 流式进度）。"""
    if not app_state.is_document_loaded:
        raise HTTPException(status_code=400, detail="请先上传并解析论文文档")

    doc_info = app_state.document_info or {}
    file_url = doc_info.get("file_url", "")
    if not file_url:
        raise HTTPException(status_code=400, detail="未找到已上传的 PDF 文件")

    pdf_filename = file_url.split("/")[-1]
    input_path = os.path.join(UPLOAD_DIR, pdf_filename)

    if not os.path.exists(input_path):
        raise HTTPException(status_code=404, detail="PDF 文件不存在")

    base_name = os.path.splitext(pdf_filename)[0]
    translated_dir = os.path.join(UPLOAD_DIR, f"{base_name}_translated")

    existing = _find_translated_files(translated_dir, base_name)
    if existing:
        async def cached():
            yield f"data: {json.dumps({'stage': 'done', 'message': '翻译已完成（缓存）', **existing}, ensure_ascii=False)}\n\n"
        return StreamingResponse(cached(), media_type="text/event-stream",
                                 headers={"Cache-Control": "no-cache", "Connection": "keep-alive"})

    from config import DEEPSEEK_API_KEY

    async def generate():
        import subprocess

        os.makedirs(translated_dir, exist_ok=True)

        yield f"data: {json.dumps({'stage': 'translating', 'message': '正在启动 PDF 翻译引擎...'}, ensure_ascii=False)}\n\n"
        await asyncio.sleep(0.1)

        translate_script = (
            "import os; "
            f"os.environ['HF_ENDPOINT']='https://hf-mirror.com'; "
            "from pdf2zh import translate; "
            "from pdf2zh.doclayout import OnnxModel; "
            "model = OnnxModel.load_available(); "
            f"translate(files=[r'{input_path}'], "
            f"lang_in='en', lang_out='zh', service='openai', "
            f"output=r'{translated_dir}', thread=2, model=model, "
            f"envs={{'OPENAI_BASE_URL': 'https://api.deepseek.com', "
            f"'OPENAI_API_KEY': '{DEEPSEEK_API_KEY}', "
            f"'OPENAI_MODEL': 'deepseek-chat'}})"
        )
        cmd = [sys.executable, "-c", translate_script]

        env = {
            **os.environ,
            "OPENAI_BASE_URL": "https://api.deepseek.com",
            "OPENAI_API_KEY": DEEPSEEK_API_KEY,
            "OPENAI_MODEL": "deepseek-chat",
            "HF_ENDPOINT": "https://hf-mirror.com",
        }

        yield f"data: {json.dumps({'stage': 'translating', 'message': '正在翻译论文（可能需要 2-5 分钟）...'}, ensure_ascii=False)}\n\n"
        await asyncio.sleep(0.1)

        try:
            process = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: subprocess.run(
                    cmd, env=env, capture_output=True, text=True, timeout=600
                )
            )

            if process.returncode != 0:
                error_msg = process.stderr[:500] if process.stderr else "未知错误"
                logger.error(f"pdf2zh 翻译失败: {error_msg}")
                yield f"data: {json.dumps({'stage': 'error', 'message': f'翻译失败: {error_msg}'}, ensure_ascii=False)}\n\n"
                return

            result = _find_translated_files(translated_dir, base_name)
            if not result:
                yield f"data: {json.dumps({'stage': 'error', 'message': '翻译完成但未找到输出文件'}, ensure_ascii=False)}\n\n"
                return

            yield f"data: {json.dumps({'stage': 'done', 'message': '翻译完成！', **result}, ensure_ascii=False)}\n\n"

        except subprocess.TimeoutExpired:
            yield f"data: {json.dumps({'stage': 'error', 'message': '翻译超时（超过10分钟），请稍后重试'}, ensure_ascii=False)}\n\n"
        except Exception as e:
            logger.exception("翻译异常")
            yield f"data: {json.dumps({'stage': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


def _find_translated_files(translated_dir: str, base_name: str) -> dict:
    """查找翻译后的 PDF 文件，返回 URL dict 或空 dict。"""
    if not os.path.isdir(translated_dir):
        return {}

    result = {}
    for f in os.listdir(translated_dir):
        if f.endswith(".pdf"):
            url = f"/api/uploads/{base_name}_translated/{f}"
            fl = f.lower()
            if "dual" in fl or "bilingual" in fl:
                result["dual_pdf_url"] = url
            elif "mono" in fl or "zh" in fl or "translated" in fl:
                result["mono_pdf_url"] = url
            else:
                if "mono_pdf_url" not in result:
                    result["mono_pdf_url"] = url
    return result


# ===================== 论文对比 =====================

@router.post("/api/compare/stream")
async def compare_stream(request: CompareRequest):
    """流式对比分析多篇论文。"""
    doc_ids = request.doc_ids
    if len(doc_ids) < 2:
        raise HTTPException(status_code=400, detail="至少选择 2 篇论文进行对比")
    if len(doc_ids) > 3:
        raise HTTPException(status_code=400, detail="最多支持 3 篇论文对比")

    papers = []
    for did in doc_ids:
        if did not in app_state.documents:
            raise HTTPException(status_code=404, detail=f"文档 {did} 不存在")
        doc_data = app_state.documents[did]
        summary = doc_data.get("summary", "")
        if not summary:
            info = doc_data.get("document_info", {})
            raise HTTPException(status_code=400, detail=f"文档「{info.get('filename', did)}」尚未分析")
        info = doc_data.get("document_info", {})
        papers.append({
            "title": info.get("title") or info.get("filename", f"论文{len(papers)+1}"),
            "summary": summary[:6000],
        })

    papers_text = ""
    for i, p in enumerate(papers):
        papers_text += f"\n{'='*50}\n论文 {i+1}: {p['title']}\n{'='*50}\n{p['summary']}\n"

    col3 = " 论文3 |" if len(papers) > 2 else ""
    sep3 = "-------|" if len(papers) > 2 else ""
    prompt = f"""你是资深学术研究员，请对以下 {len(papers)} 篇论文进行深度对比分析。
{papers_text}

请按以下结构输出（Markdown 格式）：

## 📊 对比概览表

| 维度 | 论文1 | 论文2 |{col3}
|------|-------|-------|{sep3}
| 研究目标 | | |
| 核心方法 | | |
| 数据集 | | |
| 主要结果 | | |
| 创新点 | | |

## 🔬 方法对比
详细对比各论文的技术路线差异。

## 📈 实验与结果对比
对比实验设计、评估指标和结果。

## 💡 创新点与贡献对比
分析各论文的独特贡献。

## ⚠️ 局限性对比
各论文的不足和局限。

## 🎯 总结与建议
综合评价异同，给出阅读建议。

请使用中文，对专业术语附英文原文。"""

    coordinator = get_coordinator()

    async def generate():
        try:
            full_text = ""
            for chunk in coordinator.llm_service.stream_chat(prompt):
                full_text += chunk
                yield f"data: {json.dumps({'stage': 'analyzing', 'chunk': chunk}, ensure_ascii=False)}\n\n"
                await asyncio.sleep(0.01)
            yield f"data: {json.dumps({'stage': 'done', 'analysis': full_text}, ensure_ascii=False)}\n\n"
        except Exception as e:
            logger.exception("对比分析失败")
            yield f"data: {json.dumps({'stage': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "Connection": "keep-alive"})


# ===================== 课题组讨论 =====================

@router.post("/api/lab/discuss/stream")
async def lab_discuss_stream(request: LabDiscussRequest):
    """流式课题组讨论（SSE，替代 WebSocket）。"""
    if not app_state.is_document_loaded:
        raise HTTPException(status_code=400, detail="请先上传并解析论文文档")

    coordinator = get_coordinator()
    paper_summary = app_state.current_summary or ""
    paper_title = (app_state.document_info or {}).get("title", "")
    if not paper_summary:
        raise HTTPException(status_code=400, detail="请先分析文档（需要分析结果作为讨论基础）")

    session = LabSession(
        llm_service=coordinator.llm_service,
        vector_store=coordinator.vector_store,
        paper_summary=paper_summary,
        paper_title=paper_title,
        mode=request.mode,
        user_focus=request.user_focus,
    )
    set_lab_session_cache(session)

    async def generate():
        try:
            for event in session.run_discussion_stream():
                yield f"data: {json.dumps(event, ensure_ascii=False, default=str)}\n\n"
                await asyncio.sleep(0.01)
            yield f"data: {json.dumps({'type': 'done', 'proposal': session.proposal}, ensure_ascii=False)}\n\n"
        except Exception as e:
            logger.exception("课题组讨论失败")
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


@router.get("/api/lab/proposal")
def get_lab_proposal():
    """获取最近一次课题组讨论的研究提案。"""
    session = get_lab_session_cache()
    if session and session.proposal:
        return {
            "success": True,
            "proposal": session.proposal,
            "discussion": session.shared_memory.get_discussion_text(max_chars=10000),
        }
    return {"success": False, "message": "暂无课题组讨论结果"}
