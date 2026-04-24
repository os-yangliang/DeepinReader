"""
论文阅读多智能体系统 - FastAPI 后端 API（单用户版）
"""
from fastapi import FastAPI, UploadFile, File, HTTPException, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
import uuid
import sys
import json
import asyncio
import logging

# 添加项目根目录到 Python 路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from agents.coordinator import PaperReaderCoordinator
from agents.lab.lab_session import LabSession
from services.history_store import HistoryStoreService
from services.object_indexer import ObjectIndexer
from services.paper_schema import PaperProfile
from config import SUPPORTED_EXTENSIONS, MAX_FILE_SIZE_MB, CORS_ALLOW_ORIGINS

logger = logging.getLogger(__name__)

from contextlib import asynccontextmanager
import time as _time

# 确保上传目录存在
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


@asynccontextmanager
async def lifespan(app):
    """应用生命周期管理"""
    # === 启动时 ===
    # 清理超过 7 天的上传文件
    max_age_seconds = 7 * 24 * 3600
    now = _time.time()
    cleaned = 0
    for filename in os.listdir(UPLOAD_DIR):
        filepath = os.path.join(UPLOAD_DIR, filename)
        if os.path.isfile(filepath):
            if now - os.path.getmtime(filepath) > max_age_seconds:
                try:
                    os.remove(filepath)
                    cleaned += 1
                except OSError:
                    pass
    if cleaned > 0:
        logger.info(f"已清理 {cleaned} 个过期上传文件")

    yield  # 应用运行中

    # === 关闭时 ===
    logger.info("应用关闭")


app = FastAPI(
    title="DeepinReader API",
    description="基于 LangChain + LangGraph 构建的智能论文分析与问答系统",
    version="2.0.0",
    lifespan=lifespan,
)

# 配置 CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ALLOW_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 挂载静态文件目录，用于访问上传的 PDF
app.mount("/api/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")


# ==================== 全局单例 ====================

# 应用状态（单用户，纯内存）
class AppState:
    """单用户应用状态（支持多文档）"""

    def __init__(self):
        self.current_document_id: Optional[str] = None
        self.is_document_loaded: bool = False
        self.current_summary: str = ""
        self.current_structure: str = ""
        self.document_info: dict = {}
        self.current_history_id: Optional[str] = None
        # 多文档支持
        self.documents: Dict[str, dict] = {}  # doc_id -> {document_info, summary, file_url, ...}

    def add_document(self, doc_id: str, doc_info: dict, summary: str = ""):
        """添加文档并设为活跃"""
        self.documents[doc_id] = {
            "document_info": doc_info,
            "summary": summary,
        }
        # 设为当前活跃文档
        self.current_document_id = doc_id
        self.is_document_loaded = True
        self.document_info = doc_info
        self.current_summary = summary
        self.current_structure = ""
        self.current_history_id = None

    def switch_to(self, doc_id: str) -> bool:
        """切换活跃文档"""
        if doc_id not in self.documents:
            return False
        doc = self.documents[doc_id]
        self.current_document_id = doc_id
        self.is_document_loaded = True
        self.document_info = doc["document_info"]
        self.current_summary = doc.get("summary", "")
        self.current_structure = ""
        self.current_history_id = None
        return True

    def remove_document(self, doc_id: str):
        """移除文档"""
        if doc_id in self.documents:
            del self.documents[doc_id]
        if self.current_document_id == doc_id:
            if self.documents:
                # 切换到剩余的第一个文档
                next_id = next(iter(self.documents))
                self.switch_to(next_id)
            else:
                self.clear()

    def save_summary(self, doc_id: str, summary: str):
        """保存分析结果到文档"""
        if doc_id in self.documents:
            self.documents[doc_id]["summary"] = summary
        if self.current_document_id == doc_id:
            self.current_summary = summary

    def clear(self):
        """重置所有状态"""
        self.current_document_id = None
        self.is_document_loaded = False
        self.current_summary = ""
        self.current_structure = ""
        self.document_info = {}
        self.current_history_id = None


app_state = AppState()

# Coordinator 单例缓存（避免每次请求都重新初始化 LLM + Embedding 模型）
_coordinator_instance: Optional[PaperReaderCoordinator] = None


def get_coordinator(require_llm: bool = True) -> PaperReaderCoordinator:
    """获取全局 Coordinator 单例"""
    global _coordinator_instance
    if _coordinator_instance is None:
        init_mode = "加载 LLM + Embedding 模型" if require_llm else "仅加载解析/索引能力"
        logger.info(f"首次初始化 Coordinator（{init_mode}）...")
        _coordinator_instance = PaperReaderCoordinator(require_llm=require_llm)
        logger.info("Coordinator 初始化完成")
    elif require_llm and _coordinator_instance.llm_service is None:
        logger.info("升级 Coordinator 到完整 LLM 模式...")
        _coordinator_instance = PaperReaderCoordinator(require_llm=True, vector_store=_coordinator_instance.vector_store)
        logger.info("Coordinator 已升级到完整模式")
    return _coordinator_instance


# HistoryStoreService 单例缓存
_history_store_instance: Optional[HistoryStoreService] = None


def get_history_store() -> HistoryStoreService:
    """获取全局 HistoryStoreService 单例"""
    global _history_store_instance
    if _history_store_instance is None:
        _history_store_instance = HistoryStoreService()
    return _history_store_instance


# ==================== 请求/响应模型 ====================

class ChatRequest(BaseModel):
    message: str


class ChatResponse(BaseModel):
    success: bool
    answer: str
    source_chunks: List[str] = []
    route_type: str = "general"
    confidence: float = 0.0
    warnings: List[str] = []
    evidence_summary: List[str] = []
    reasoning_trace: List[str] = []
    reasoning_paths: List[List[str]] = []
    claim_nodes: List[str] = []
    evidence_nodes: List[str] = []
    result_nodes: List[str] = []


class AnalysisResponse(BaseModel):
    success: bool
    status: str
    document_info: dict = {}
    structure: str = ""
    summary: str = ""
    error: str = ""


class DocumentInfoResponse(BaseModel):
    is_loaded: bool
    info: dict = {}
    structure: str = ""
    summary: str = ""


class HistoryItem(BaseModel):
    id: str
    filename: str
    title: str = ""
    file_type: str = ""
    page_count: int = 0
    word_count: int = 0
    processing_time: float = 0
    analyzed_at: str = ""


class HistoryListResponse(BaseModel):
    history: List[HistoryItem] = []
    current_id: Optional[str] = None


class ChatHistoryItem(BaseModel):
    id: str
    role: str
    content: str
    timestamp: str = ""
    source_chunks: List[str] = []
    route_type: str = "general"
    confidence: float = 0.0
    warnings: List[str] = []
    evidence_summary: List[str] = []
    reasoning_trace: List[str] = []


class ProfileSummaryResponse(BaseModel):
    success: bool
    document_id: str = ""
    title: str = ""
    abstract: str = ""
    counts: Dict[str, int] = {}
    contributions: List[str] = []
    limitations: List[str] = []
    keywords: List[str] = []


class ProfileDetailResponse(ProfileSummaryResponse):
    sections: List[dict] = []
    claims: List[dict] = []
    evidences: List[dict] = []
    experiments: List[dict] = []
    results: List[dict] = []
    graph: Dict[str, List[str]] = {}


def _load_profile_for_document(document_id: str) -> Optional[PaperProfile]:
    if not document_id:
        return None
    coordinator = get_coordinator()
    return ObjectIndexer(coordinator.vector_store).load_profile(document_id)


def _build_profile_summary(profile: PaperProfile) -> Dict[str, Any]:
    return {
        "success": True,
        "document_id": profile.document_id,
        "title": profile.title,
        "abstract": profile.abstract[:1200],
        "counts": {
            "sections": len(profile.sections),
            "claims": len(profile.claims),
            "evidences": len(profile.evidences),
            "experiments": len(profile.experiments),
            "results": len(profile.results),
        },
        "contributions": profile.contributions[:8],
        "limitations": profile.limitations[:8],
        "keywords": profile.keywords[:12],
    }


def _build_profile_detail(profile: PaperProfile) -> Dict[str, Any]:
    payload = _build_profile_summary(profile)
    payload.update({
        "sections": [section.model_dump() for section in profile.sections],
        "claims": [claim.model_dump() for claim in profile.claims],
        "evidences": [evidence.model_dump() for evidence in profile.evidences],
        "experiments": [experiment.model_dump() for experiment in profile.experiments],
        "results": [result.model_dump() for result in profile.results],
        "graph": profile.graph,
    })
    return payload


# ==================== WebSocket 统一流式端点 ====================

# 活跃的 WS 任务（用于取消）
_active_ws_tasks: Dict[str, asyncio.Task] = {}
_cancel_flags: Dict[str, bool] = {}


@app.websocket("/ws")
async def websocket_endpoint(ws: WebSocket):
    """统一 WebSocket 端点 — 多路复用所有流式操作"""
    await ws.accept()

    async def send(msg_type: str, request_id: str, data: dict):
        """统一发送格式"""
        try:
            await ws.send_json({"type": msg_type, "request_id": request_id, "data": data})
        except Exception:
            pass

    async def handle_analyze(request_id: str, data: dict):
        """流式分析"""
        if not app_state.is_document_loaded:
            await send("error", request_id, {"message": "请先上传文档"})
            return
        coordinator = get_coordinator()
        try:
            full_analysis = ""
            for event in coordinator.stream_analysis():
                if _cancel_flags.get(request_id):
                    await send("cancelled", request_id, {})
                    return
                stage = event.get("stage", "")
                if stage == "analyzing" and event.get("chunk"):
                    full_analysis += event["chunk"]
                if stage == "done":
                    analysis_text = event.get("analysis", "")
                    app_state.current_summary = analysis_text
                    if app_state.current_document_id:
                        app_state.save_summary(app_state.current_document_id, analysis_text)
                    try:
                        doc_info = app_state.document_info or {}
                        history_id = get_history_store().add_analysis_history(
                            document_id=app_state.current_document_id or "",
                            filename=doc_info.get("filename", ""),
                            title=doc_info.get("title", ""),
                            file_type=doc_info.get("file_type", ""),
                            page_count=doc_info.get("page_count", 0),
                            word_count=doc_info.get("word_count", 0),
                            processing_time=0, structure="",
                            summary=analysis_text,
                        )
                        app_state.current_history_id = history_id
                    except Exception as e:
                        logger.error(f"保存历史失败: {e}")
                await send("stream", request_id, event)
                await asyncio.sleep(0.01)
            await send("done", request_id, {})
        except Exception as e:
            await send("error", request_id, {"message": str(e)})

    async def handle_chat(request_id: str, data: dict):
        """流式问答"""
        message = data.get("message", "").strip()
        if not message:
            await send("error", request_id, {"message": "消息不能为空"})
            return
        if not app_state.is_document_loaded:
            await send("error", request_id, {"message": "请先上传文档"})
            return
        coordinator = get_coordinator()
        try:
            if app_state.current_document_id:
                get_history_store().add_chat_message(
                    document_id=app_state.current_document_id,
                    role="user",
                    content=message,
                    route_type="general",
                    confidence=0.0,
                    warnings=[],
                    evidence_summary=[],
                    reasoning_trace=[],
                )
        except Exception:
            pass
        full_response = ""
        result = None
        try:
            result = coordinator.ask_question(message)
            if not result.success:
                await send("error", request_id, {"message": result.error_message or "问答失败"})
                return
            full_response = result.answer
            for chunk in coordinator.ask_question_stream(message):
                if _cancel_flags.get(request_id):
                    await send("cancelled", request_id, {"partial": full_response})
                    return
                await send("stream", request_id, {"chunk": chunk})
                await asyncio.sleep(0.01)
            try:
                if app_state.current_document_id:
                    get_history_store().add_chat_message(
                        document_id=app_state.current_document_id,
                        role="assistant",
                        content=full_response,
                        source_chunks=result.source_chunks[:3] if result and result.source_chunks else [],
                        route_type=result.route_type if result else "general",
                        confidence=result.confidence if result else 0.0,
                        warnings=result.warnings if result else [],
                        evidence_summary=result.evidence_summary if result else [],
                        reasoning_trace=result.reasoning_trace if result else [],
                    )
            except Exception:
                pass
            await send("done", request_id, {
                "full_response": full_response,
                "route_type": result.route_type if result else "general",
                "confidence": result.confidence if result else 0.0,
                "warnings": result.warnings if result else [],
                "evidence_summary": result.evidence_summary if result else [],
                "source_chunks": result.source_chunks[:3] if result and result.source_chunks else [],
            })
        except Exception as e:
            await send("error", request_id, {"message": str(e)})

    async def handle_translate(request_id: str, data: dict):
        """流式翻译 — 翻译使用 pdf2zh 子进程，不走 WS，强制 SSE 降级"""
        await send("error", request_id, {"message": "translate_use_sse"})

    async def handle_code_generate(request_id: str, data: dict):
        """流式代码生成"""
        user_request = data.get("user_request", "")
        target_framework = data.get("target_framework", "Python (PyTorch)")
        if not app_state.is_document_loaded:
            await send("error", request_id, {"message": "请先上传文档"})
            return
        coordinator = get_coordinator()
        try:
            coordinator.current_state["summary"] = app_state.current_summary
            for chunk in coordinator.generate_code_stream(
                user_request=user_request,
                target_framework=target_framework,
                paper_summary=app_state.current_summary[:500],
            ):
                if _cancel_flags.get(request_id):
                    await send("cancelled", request_id, {})
                    return
                await send("stream", request_id, {"chunk": chunk})
                await asyncio.sleep(0.01)
            await send("done", request_id, {})
        except Exception as e:
            await send("error", request_id, {"message": str(e)})

    async def handle_compare(request_id: str, data: dict):
        """流式对比分析"""
        doc_ids = data.get("doc_ids", [])
        if len(doc_ids) < 2 or len(doc_ids) > 3:
            await send("error", request_id, {"message": "请选择 2-3 篇论文"})
            return
        papers = []
        for did in doc_ids:
            if did not in app_state.documents:
                await send("error", request_id, {"message": f"文档 {did} 不存在"})
                return
            doc_data = app_state.documents[did]
            summary = doc_data.get("summary", "")
            if not summary:
                info = doc_data.get("document_info", {})
                await send("error", request_id, {"message": f"文档「{info.get('filename', did)}」尚未分析"})
                return
            info = doc_data.get("document_info", {})
            papers.append({
                "title": info.get("title") or info.get("filename", f"论文{len(papers)+1}"),
                "summary": summary[:6000],
            })
        papers_text = ""
        for i, p in enumerate(papers):
            papers_text += f"\n{'='*50}\n论文 {i+1}: {p['title']}\n{'='*50}\n{p['summary']}\n"
        col3 = " 论文3 |" if len(papers) > 2 else ""
        sep3 = "-------|" if len(papers) > 2 else ""
        prompt = f"""你是资深学术研究员，请对以下 {len(papers)} 篇论文进行深度对比分析。
{papers_text}
请按以下结构输出（Markdown 格式）：
## 📊 对比概览表
| 维度 | 论文1 | 论文2 |{col3}
|------|-------|-------|{sep3}
| 研究目标 | | |
| 核心方法 | | |
| 数据集 | | |
| 主要结果 | | |
| 创新点 | | |
## 🔬 方法对比
## 📈 实验与结果对比
## 💡 创新点与贡献对比
## ⚠️ 局限性对比
## 🎯 总结与建议
请使用中文，对专业术语附英文原文。"""
        coordinator = get_coordinator()
        try:
            full_text = ""
            for chunk in coordinator.llm_service.stream_chat(prompt):
                if _cancel_flags.get(request_id):
                    await send("cancelled", request_id, {})
                    return
                full_text += chunk
                await send("stream", request_id, {"stage": "analyzing", "chunk": chunk})
                await asyncio.sleep(0.01)
            await send("done", request_id, {"stage": "done", "analysis": full_text})
        except Exception as e:
            await send("error", request_id, {"message": str(e)})

    async def handle_lab_discuss(request_id: str, data: dict):
        """流式课题组讨论"""
        if not app_state.is_document_loaded:
            await send("error", request_id, {"message": "请先上传文档"})
            return
        coordinator = get_coordinator()
        paper_summary = app_state.current_summary or ""
        paper_title = (app_state.document_info or {}).get("title", "")
        if not paper_summary:
            await send("error", request_id, {"message": "请先分析文档（需要分析结果作为讨论基础）"})
            return
        mode = data.get("mode", "quick")
        user_focus = data.get("user_focus", "")
        session = LabSession(
            llm_service=coordinator.llm_service,
            vector_store=coordinator.vector_store,
            paper_summary=paper_summary,
            paper_title=paper_title,
            mode=mode,
            user_focus=user_focus,
        )
        try:
            for event in session.run_discussion_stream():
                if _cancel_flags.get(request_id):
                    await send("cancelled", request_id, {})
                    return
                await send("stream", request_id, event)
                await asyncio.sleep(0.01)
            await send("done", request_id, {
                "proposal": session.proposal,
            })
        except Exception as e:
            logger.exception("课题组讨论失败")
            await send("error", request_id, {"message": str(e)})

    # 路由表
    handlers = {
        "analyze": handle_analyze,
        "chat": handle_chat,
        "translate": handle_translate,
        "code_generate": handle_code_generate,
        "compare": handle_compare,
        "lab_discuss": handle_lab_discuss,
    }

    try:
        while True:
            msg = await ws.receive_json()
            msg_type = msg.get("type", "")
            request_id = msg.get("request_id", "")

            if msg_type == "cancel":
                # 取消指定任务
                _cancel_flags[request_id] = True
                task = _active_ws_tasks.get(request_id)
                if task and not task.done():
                    task.cancel()
                await send("cancelled", request_id, {})
                continue

            if msg_type == "ping":
                await send("pong", request_id, {})
                continue

            handler = handlers.get(msg_type)
            if not handler:
                await send("error", request_id, {"message": f"未知消息类型: {msg_type}"})
                continue

            # 启动后台任务
            _cancel_flags[request_id] = False
            task = asyncio.create_task(handler(request_id, msg.get("data", {})))
            _active_ws_tasks[request_id] = task

            async def cleanup(rid, t):
                try:
                    await t
                except asyncio.CancelledError:
                    pass
                finally:
                    _active_ws_tasks.pop(rid, None)
                    _cancel_flags.pop(rid, None)

            asyncio.create_task(cleanup(request_id, task))

    except WebSocketDisconnect:
        logger.info("WebSocket 客户端断开")
    except Exception as e:
        logger.error(f"WebSocket 错误: {e}")
    finally:
        # 清理所有活跃任务
        for rid, task in list(_active_ws_tasks.items()):
            _cancel_flags[rid] = True
            task.cancel()
        _active_ws_tasks.clear()
        _cancel_flags.clear()


# ==================== API 端点 ====================

@app.get("/")
async def root():
    return {"message": "论文阅读助手 API", "version": "2.0.0"}


@app.get("/api/status")
async def get_status():
    """获取系统状态"""
    return {
        "is_document_loaded": app_state.is_document_loaded,
        "has_coordinator": True,
    }


@app.get("/api/document", response_model=DocumentInfoResponse)
async def get_document_info():
    """获取当前文档信息"""
    return DocumentInfoResponse(
        is_loaded=app_state.is_document_loaded,
        info=app_state.document_info,
        structure=app_state.current_structure,
        summary=app_state.current_summary,
    )


@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...)):
    """上传文档：保存 + 解析 + 建立向量索引（无 LLM，~5秒）"""
    coordinator = get_coordinator(require_llm=False)

    try:
        filename = file.filename
        _, ext = os.path.splitext(filename)

        if ext.lower() not in SUPPORTED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件格式: {ext}，请上传 PDF 或 Word 文档",
            )

        file_bytes = await file.read()
        max_bytes = int(MAX_FILE_SIZE_MB * 1024 * 1024)
        if len(file_bytes) > max_bytes:
            raise HTTPException(
                status_code=413,
                detail=f"文件过大：{len(file_bytes) / (1024 * 1024):.2f}MB，最大支持 {MAX_FILE_SIZE_MB}MB",
            )

        # 保存文件
        file_id = str(uuid.uuid4())
        save_filename = f"{file_id}{ext}"
        save_path = os.path.join(UPLOAD_DIR, save_filename)
        
        def _save_and_index():
            with open(save_path, "wb") as f:
                f.write(file_bytes)
            return coordinator.parse_and_index(file_bytes, filename)
            
        doc_info = await asyncio.to_thread(_save_and_index)
        doc_info["file_url"] = f"/api/uploads/{save_filename}"

        # 添加到多文档列表并设为活跃
        doc_id = doc_info.get("document_id", file_id)
        app_state.add_document(doc_id, doc_info)

        if coordinator.current_state is not None:
            coordinator.current_state["document_id"] = doc_id
            coordinator.current_state["parsed_doc"] = coordinator.current_state.get("parsed_doc")
            coordinator.current_state["summary"] = ""
            coordinator.current_state["structure_info"] = doc_info.get("structure_info", "")
            coordinator.current_state["keywords"] = ""
            coordinator.current_state["current_stage"] = "indexed"
        if coordinator.qa_agent:
            coordinator.qa_agent.set_document_context(
                doc_id=doc_id,
                paper_title=doc_info.get("title", ""),
                paper_summary="",
            )

        return {
            "success": True,
            "message": f"文档就绪：{doc_info.get('page_count', 0)} 页，{doc_info.get('word_count', 0)} 字",
            "document_info": doc_info,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.exception("文档上传解析失败")
        return {"success": False, "error": str(e)}


@app.post("/api/analyze/stream")
async def analyze_stream():
    """流式 AI 分析（独立于上传，需要先上传文档）"""
    if not app_state.is_document_loaded:
        raise HTTPException(status_code=400, detail="请先上传文档")

    coordinator = get_coordinator()

    async def generate():
        try:
            full_analysis = ""
            for event in coordinator.stream_analysis():
                stage = event.get("stage", "")

                if stage == "analyzing" and event.get("chunk"):
                    full_analysis += event["chunk"]

                if stage == "done":
                    # 更新应用状态（同时存入多文档注册表）
                    analysis_text = event.get("analysis", "")
                    app_state.current_summary = analysis_text
                    if app_state.current_document_id:
                        app_state.save_summary(app_state.current_document_id, analysis_text)

                    # 保存历史记录
                    try:
                        doc_info = app_state.document_info or {}
                        history_id = get_history_store().add_analysis_history(
                            document_id=app_state.current_document_id or "",
                            filename=doc_info.get("filename", ""),
                            title=doc_info.get("title", ""),
                            file_type=doc_info.get("file_type", ""),
                            page_count=doc_info.get("page_count", 0),
                            word_count=doc_info.get("word_count", 0),
                            processing_time=0,
                            structure="",
                            summary=event.get("analysis", ""),
                        )
                        app_state.current_history_id = history_id
                    except Exception as e:
                        logger.error(f"保存历史记录失败: {e}")

                yield f"data: {json.dumps(event, ensure_ascii=False, default=str)}\n\n"
                await asyncio.sleep(0.01)

        except Exception as e:
            logger.exception("流式分析失败")
            yield f"data: {json.dumps({'stage': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


@app.post("/api/upload/stream")
async def upload_and_analyze_stream(file: UploadFile = File(...)):
    """流式上传并分析论文（SSE）"""
    coordinator = get_coordinator()

    # 验证文件
    filename = file.filename
    _, ext = os.path.splitext(filename)
    if ext.lower() not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {ext}，请上传 PDF 或 Word 文档",
        )

    file_bytes = await file.read()
    max_bytes = int(MAX_FILE_SIZE_MB * 1024 * 1024)
    if len(file_bytes) > max_bytes:
        raise HTTPException(
            status_code=413,
            detail=f"文件过大：{len(file_bytes) / (1024 * 1024):.2f}MB，最大支持 {MAX_FILE_SIZE_MB}MB",
        )

    # 保存文件（异步写入避免阻塞事件循环）
    file_id = str(uuid.uuid4())
    save_filename = f"{file_id}{ext}"
    save_path = os.path.join(UPLOAD_DIR, save_filename)
    await asyncio.to_thread(lambda: open(save_path, "wb").write(file_bytes))
    file_url = f"/api/uploads/{save_filename}"

    async def generate():
        try:
            for event in coordinator.process_document_stream(file_bytes, filename):
                stage = event.get("stage", "")

                # 在 done 阶段更新应用状态和历史记录
                if stage == "done":
                    doc_info = event.get("document_info", {})
                    doc_info["file_url"] = file_url

                    app_state.is_document_loaded = True
                    app_state.current_document_id = doc_info.get("document_id")
                    app_state.current_summary = event.get("analysis", "")
                    app_state.current_structure = ""
                    app_state.document_info = doc_info

                    # 保存历史记录
                    try:
                        history_id = get_history_store().add_analysis_history(
                            document_id=app_state.current_document_id,
                            filename=doc_info.get("filename", ""),
                            title=doc_info.get("title", ""),
                            file_type=doc_info.get("file_type", ""),
                            page_count=doc_info.get("page_count", 0),
                            word_count=doc_info.get("word_count", 0),
                            processing_time=doc_info.get("processing_time", 0),
                            structure="",
                            summary=event.get("analysis", ""),
                        )
                        app_state.current_history_id = history_id
                    except Exception as e:
                        logger.error(f"保存历史记录失败: {e}")

                    event["document_info"] = doc_info

                # 如果有 file_url 需要在 parsed 阶段也带上
                if stage == "parsed":
                    event["file_url"] = file_url

                yield f"data: {json.dumps(event, ensure_ascii=False, default=str)}\n\n"
                await asyncio.sleep(0.01)

        except Exception as e:
            logger.exception("流式分析失败")
            yield f"data: {json.dumps({'stage': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


class TranslateTextRequest(BaseModel):
    text: str


@app.post("/api/translate/text")
def translate_text(request: TranslateTextRequest):
    """划词翻译 - 翻译选中的文本片段"""
    text = request.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="文本不能为空")
    if len(text) > 500:
        text = text[:500]

    coordinator = get_coordinator()
    try:
        result = coordinator.llm_service.chat_sync(
            user_message=f"请将以下英文翻译成中文，只输出翻译结果，不要解释：\n\n{text}",
            system_prompt="你是专业的学术论文翻译员。翻译要准确、专业、通顺。对于专业术语，在翻译后用括号附上英文原文。",
            chat_history=[],
        )
        return {"translation": result.strip(), "original": text}
    except Exception as e:
        logger.error(f"划词翻译失败: {e}")
        raise HTTPException(status_code=500, detail=f"翻译失败: {str(e)}")


# ===================== 导出报告 =====================

class ExportRequest(BaseModel):
    annotations: List[dict] = []  # [{text, note, color, timestamp}]

@app.post("/api/export/report")
def export_report(request: ExportRequest = ExportRequest()):
    """导出分析报告为 Word 文档"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import re, tempfile, datetime

    if not app_state.current_summary:
        raise HTTPException(status_code=400, detail="暂无分析结果可导出")

    doc_info = app_state.document_info or {}
    summary = app_state.current_summary

    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    # ---- 封面 ----
    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 30)
    run.font.color.rgb = RGBColor(14, 165, 233)
    run.font.size = Pt(14)

    # 标题
    title_text = doc_info.get("title") or doc_info.get("filename", "论文分析报告")
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(15, 23, 42)

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI 智能分析报告")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 116, 139)
    run.font.italic = True

    # 元信息
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.space_after = Pt(6)
    parts = []
    if doc_info.get("page_count"):
        parts.append(f"📄 {doc_info['page_count']} 页")
    if doc_info.get("word_count"):
        parts.append(f"📝 {doc_info['word_count']:,} 字")
    parts.append(f"🕐 {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run = meta.add_run("  |  ".join(parts))
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(148, 163, 184)

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 30)
    run.font.color.rgb = RGBColor(14, 165, 233)
    run.font.size = Pt(14)

    doc.add_page_break()

    # ---- 目录标题 ----
    toc_title = doc.add_heading("目录", level=1)
    for run in toc_title.runs:
        run.font.color.rgb = RGBColor(14, 165, 233)

    # 从 summary 提取 h2 标题作为目录
    h2_titles = re.findall(r'^##\s+(.+)$', summary, re.MULTILINE)
    if h2_titles:
        for i, t in enumerate(h2_titles):
            clean = re.sub(r'[#*`]', '', t).strip()
            p = doc.add_paragraph(f"{i+1}. {clean}", style='List Number')
            p.paragraph_format.space_after = Pt(4)
    else:
        doc.add_paragraph("（AI 自动生成分析内容）")

    doc.add_page_break()

    # ---- 正文：解析 Markdown 渲染到 docx ----
    section_title = doc.add_heading("AI 分析报告", level=1)
    for run in section_title.runs:
        run.font.color.rgb = RGBColor(14, 165, 233)

    _render_markdown_to_docx(doc, summary)

    # ---- 标注部分 ----
    if request.annotations:
        doc.add_page_break()
        anno_title = doc.add_heading("📝 阅读标注", level=1)
        for run in anno_title.runs:
            run.font.color.rgb = RGBColor(251, 191, 36)

        doc.add_paragraph(f"共 {len(request.annotations)} 条标注")

        for i, anno in enumerate(request.annotations):
            h = doc.add_heading(f"标注 {i+1}", level=3)
            # 引用文本
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f"「{anno.get('text', '')}」")
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 116, 139)
            # 笔记
            if anno.get("note"):
                note_p = doc.add_paragraph()
                note_p.paragraph_format.left_indent = Cm(1)
                run = note_p.add_run(f"📌 笔记: {anno['note']}")
                run.font.size = Pt(10)
            # 时间
            if anno.get("timestamp"):
                try:
                    ts = datetime.datetime.fromisoformat(anno["timestamp"].replace("Z", "+00:00"))
                    time_p = doc.add_paragraph()
                    time_p.paragraph_format.left_indent = Cm(1)
                    run = time_p.add_run(ts.strftime("%Y-%m-%d %H:%M"))
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(148, 163, 184)
                except:
                    pass

    # ---- 页脚 ----
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.space_before = Pt(24)
    run = footer_p.add_run("— 由 PaperReader 智能论文阅读助手生成 —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(148, 163, 184)
    run.font.italic = True

    # 保存
    filename = (doc_info.get("filename", "report") or "report").replace(".pdf", "").replace(".docx", "")
    output_path = os.path.join(tempfile.gettempdir(), f"{filename}_analysis_report.docx")
    doc.save(output_path)

    from fastapi.responses import FileResponse
    return FileResponse(
        path=output_path,
        filename=f"{filename}_analysis_report.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _render_markdown_to_docx(doc, md_text: str):
    """将 Markdown 文本渲染到 Word 文档"""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

    lines = md_text.split("\n")
    in_code_block = False
    code_lang = ""
    code_lines = []

    for line in lines:
        # 代码块
        if line.strip().startswith("```"):
            if in_code_block:
                # 结束代码块
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 116, 139)
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lang = line.strip().replace("```", "").strip()
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        stripped = line.strip()
        if not stripped:
            continue

        # 标题
        if stripped.startswith("#### "):
            h = doc.add_heading(re.sub(r'[#]', '', stripped[5:]).strip(), level=4)
            h.paragraph_format.space_before = Pt(8)
        elif stripped.startswith("### "):
            h = doc.add_heading(re.sub(r'[#]', '', stripped[4:]).strip(), level=3)
            h.paragraph_format.space_before = Pt(10)
        elif stripped.startswith("## "):
            h = doc.add_heading(re.sub(r'[#]', '', stripped[3:]).strip(), level=2)
            h.paragraph_format.space_before = Pt(14)
        elif stripped.startswith("# "):
            h = doc.add_heading(re.sub(r'[#]', '', stripped[2:]).strip(), level=1)
            h.paragraph_format.space_before = Pt(16)
        # 引用
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            content = stripped[2:]
            run = p.add_run(content)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 116, 139)
        # 无序列表
        elif stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, content)
        # 有序列表
        elif re.match(r'^\d+\.\s', stripped):
            content = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style="List Number")
            _add_rich_text(p, content)
        # 分隔线
        elif stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 40)
            run.font.color.rgb = RGBColor(203, 213, 225)
            run.font.size = Pt(8)
        # 普通段落
        else:
            p = doc.add_paragraph()
            _add_rich_text(p, stripped)


def _add_rich_text(paragraph, text: str):
    """解析 Markdown inline 格式（粗体、斜体、行内代码）"""
    from docx.shared import Pt, RGBColor
    import re

    # 分割: **bold**, *italic*, `code`
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(14, 165, 233)
        else:
            paragraph.add_run(part)


@app.post("/api/chat", response_model=ChatResponse)
def chat(request: ChatRequest):
    """聊天问答"""
    coordinator = get_coordinator()

    if not request.message.strip():
        raise HTTPException(status_code=400, detail="消息不能为空")

    if not app_state.is_document_loaded:
        raise HTTPException(status_code=400, detail="请先上传并解析论文文档")

    # 保存用户消息到历史
    try:
        if app_state.current_document_id:
            get_history_store().add_chat_message(
                document_id=app_state.current_document_id,
                role="user",
                content=request.message,
                route_type="general",
                confidence=0.0,
                warnings=[],
                evidence_summary=[],
                reasoning_trace=[],
            )
    except Exception as e:
        logger.error(f"保存消息失败: {e}")

    # 获取回答
    result = coordinator.ask_question(request.message)

    if result.success:
        # 保存助手回复到历史
        try:
            if app_state.current_document_id:
                get_history_store().add_chat_message(
                    document_id=app_state.current_document_id,
                    role="assistant",
                    content=result.answer,
                    source_chunks=result.source_chunks[:3] if result.source_chunks else [],
                    route_type=result.route_type,
                    confidence=result.confidence,
                    warnings=result.warnings,
                    evidence_summary=result.evidence_summary,
                    reasoning_trace=result.reasoning_trace,
                )
        except Exception as e:
            logger.error(f"保存回复失败: {e}")

        return ChatResponse(
            success=True,
            answer=result.answer,
            source_chunks=result.source_chunks[:3] if result.source_chunks else [],
            route_type=result.route_type,
            confidence=result.confidence,
            warnings=result.warnings,
            evidence_summary=result.evidence_summary,
            reasoning_trace=result.reasoning_trace,
            reasoning_paths=result.reasoning_paths,
            claim_nodes=result.claim_nodes,
            evidence_nodes=result.evidence_nodes,
            result_nodes=result.result_nodes,
        )
    else:
        return ChatResponse(
            success=False,
            answer=f"回答失败: {result.error_message}",
            source_chunks=[],
            route_type="general",
            confidence=0.0,
            warnings=[],
            evidence_summary=[],
            reasoning_trace=[],
            reasoning_paths=[],
            claim_nodes=[],
            evidence_nodes=[],
            result_nodes=[],
        )


@app.post("/api/chat/stream")
async def chat_stream(request: ChatRequest):
    """流式聊天问答"""
    coordinator = get_coordinator()

    if not request.message.strip():
        raise HTTPException(status_code=400, detail="消息不能为空")

    if not app_state.is_document_loaded:
        raise HTTPException(status_code=400, detail="请先上传并解析论文文档")

    # 保存用户消息
    try:
        if app_state.current_document_id:
            get_history_store().add_chat_message(
                document_id=app_state.current_document_id,
                role="user",
                content=request.message,
                route_type="general",
                confidence=0.0,
                warnings=[],
                evidence_summary=[],
                reasoning_trace=[],
            )
    except Exception as e:
        logger.error(f"保存消息失败: {e}")

    async def generate():
        full_response = ""
        try:
            result = coordinator.ask_question(request.message)
            if not result.success:
                yield f"data: {json.dumps({'error': result.error_message or '问答失败'}, ensure_ascii=False)}\n\n"
                return

            for chunk in coordinator.ask_question_stream(request.message):
                full_response += chunk
                yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
                await asyncio.sleep(0.01)

            try:
                if app_state.current_document_id:
                    get_history_store().add_chat_message(
                        document_id=app_state.current_document_id,
                        role="assistant",
                        content=result.answer,
                        source_chunks=result.source_chunks[:3] if result.source_chunks else [],
                        route_type=result.route_type,
                        confidence=result.confidence,
                        warnings=result.warnings,
                        evidence_summary=result.evidence_summary,
                        reasoning_trace=result.reasoning_trace,
                    )
            except Exception as e:
                logger.error(f"保存回复失败: {e}")

            yield f"data: {json.dumps({'done': True, 'route_type': result.route_type, 'confidence': result.confidence, 'warnings': result.warnings, 'evidence_summary': result.evidence_summary, 'reasoning_trace': result.reasoning_trace, 'reasoning_paths': result.reasoning_paths, 'claim_nodes': result.claim_nodes, 'evidence_nodes': result.evidence_nodes, 'result_nodes': result.result_nodes, 'source_chunks': result.source_chunks[:3] if result.source_chunks else []}, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


@app.get("/api/suggestions")
def get_suggestions():
    """获取建议问题"""
    if not app_state.is_document_loaded:
        return {"questions": []}

    coordinator = get_coordinator()
    questions = coordinator.get_suggested_questions()
    return {"questions": questions}


@app.post("/api/mindmap")
def generate_mindmap():
    """生成论文思维导图"""
    if not app_state.is_document_loaded:
        raise HTTPException(status_code=400, detail="请先上传并解析论文文档")

    coordinator = get_coordinator()

    paper_title = ""
    paper_summary = ""
    if coordinator.current_state:
        parsed_doc = coordinator.current_state.get("parsed_doc")
        if parsed_doc:
            paper_title = parsed_doc.title or ""
        paper_summary = coordinator.current_state.get("summary", "") or ""

    if not paper_summary:
        raise HTTPException(status_code=400, detail="请先完成论文分析")

    from prompts.templates import MINDMAP_PROMPT

    prompt = MINDMAP_PROMPT.format(
        paper_title=paper_title or "未知标题",
        paper_summary=paper_summary[:3000],
    )

    try:
        result = coordinator.llm_service.chat_sync(
            user_message=prompt,
            system_prompt="你是学术论文结构分析专家。只输出纯 JSON，不要其他文字。",
            chat_history=[],
        )
        # 提取 JSON
        import re
        # 尝试提取 ```json ... ``` 代码块
        match = re.search(r"```(?:json)?\s*([\s\S]*?)```", result)
        json_str = match.group(1).strip() if match else result.strip()
        
        # 解析 JSON
        tree_data = json.loads(json_str)
        
        return {"tree": tree_data, "title": paper_title}
    except json.JSONDecodeError as e:
        logger.error(f"思维导图 JSON 解析失败: {e}, raw: {result[:200]}")
        raise HTTPException(status_code=500, detail="AI 返回的格式不正确，请重试")
    except Exception as e:
        logger.error(f"思维导图生成失败: {e}")
        raise HTTPException(status_code=500, detail=f"生成失败: {str(e)}")


class SearchRequest(BaseModel):
    query: str = ""
    limit: int = 10


@app.post("/api/search")
async def search_papers(request: SearchRequest):
    """搜索相关论文（Semantic Scholar API）"""
    import httpx

    query = request.query.strip()

    # 如果没有提供查询，使用当前文档标题
    if not query:
        if app_state.document_info:
            query = app_state.document_info.get("title", "") or app_state.document_info.get("filename", "")
        if not query:
            raise HTTPException(status_code=400, detail="请提供搜索关键词或先上传文档")

    try:
        data = None
        async with httpx.AsyncClient(timeout=15.0) as client:
            for attempt in range(3):
                resp = await client.get(
                    "https://api.semanticscholar.org/graph/v1/paper/search",
                    params={
                        "query": query[:200],  # 限制查询长度
                        "limit": min(request.limit, 20),
                        "fields": "title,abstract,authors,year,citationCount,url,openAccessPdf,externalIds",
                    },
                )
                if resp.status_code == 429:
                    wait = 2 ** attempt
                    logger.warning(f"Semantic Scholar 429, 等待 {wait}s 重试...")
                    await asyncio.sleep(wait)
                    continue
                resp.raise_for_status()
                data = resp.json()
                break

        if data is None:
            raise HTTPException(status_code=429, detail="搜索请求过于频繁，请稍后再试")

        papers = []
        for p in data.get("data", []):
            authors = [a.get("name", "") for a in (p.get("authors") or [])[:5]]
            pdf_url = ""
            if p.get("openAccessPdf"):
                pdf_url = p["openAccessPdf"].get("url", "")
            doi = ""
            if p.get("externalIds"):
                doi = p["externalIds"].get("DOI", "")

            papers.append({
                "title": p.get("title", ""),
                "authors": authors,
                "year": p.get("year"),
                "citationCount": p.get("citationCount", 0),
                "abstract": (p.get("abstract") or "")[:300],
                "url": p.get("url", ""),
                "pdfUrl": pdf_url,
                "doi": doi,
            })

        return {"query": query, "total": data.get("total", 0), "papers": papers}

    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="搜索超时，请稍后重试")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"搜索失败: {e}")
        raise HTTPException(status_code=500, detail=f"搜索失败: {str(e)}")


@app.post("/api/translate/stream")
async def translate_stream():
    """使用 pdf2zh 翻译论文 PDF（SSE 流式进度）"""
    if not app_state.is_document_loaded:
        raise HTTPException(status_code=400, detail="请先上传并解析论文文档")

    # 找到上传的 PDF 文件路径
    doc_info = app_state.document_info or {}
    file_url = doc_info.get("file_url", "")
    if not file_url:
        raise HTTPException(status_code=400, detail="未找到已上传的 PDF 文件")

    # 从 /api/uploads/xxx.pdf 提取文件名
    pdf_filename = file_url.split("/")[-1]
    input_path = os.path.join(UPLOAD_DIR, pdf_filename)

    if not os.path.exists(input_path):
        raise HTTPException(status_code=404, detail="PDF 文件不存在")

    # 检查是否已翻译过（缓存）
    base_name = os.path.splitext(pdf_filename)[0]
    translated_dir = os.path.join(UPLOAD_DIR, f"{base_name}_translated")
    
    # 查找已有翻译结果
    existing = _find_translated_files(translated_dir, base_name)
    if existing:
        async def cached():
            yield f"data: {json.dumps({'stage': 'done', 'message': '翻译已完成（缓存）', **existing}, ensure_ascii=False)}\n\n"
        return StreamingResponse(cached(), media_type="text/event-stream",
                                 headers={"Cache-Control": "no-cache", "Connection": "keep-alive"})

    from config import DEEPSEEK_API_KEY

    async def generate():
        import subprocess

        os.makedirs(translated_dir, exist_ok=True)

        yield f"data: {json.dumps({'stage': 'translating', 'message': '正在启动 PDF 翻译引擎...'}, ensure_ascii=False)}\n\n"
        await asyncio.sleep(0.1)

        # 构建 pdf2zh 命令（需要先加载布局检测模型）
        translate_script = (
            "import os; "
            f"os.environ['HF_ENDPOINT']='https://hf-mirror.com'; "
            "from pdf2zh import translate; "
            "from pdf2zh.doclayout import OnnxModel; "
            "model = OnnxModel.load_available(); "
            f"translate(files=[r'{input_path}'], "
            f"lang_in='en', lang_out='zh', service='openai', "
            f"output=r'{translated_dir}', thread=2, model=model, "
            f"envs={{'OPENAI_BASE_URL': 'https://api.deepseek.com', "
            f"'OPENAI_API_KEY': '{DEEPSEEK_API_KEY}', "
            f"'OPENAI_MODEL': 'deepseek-chat'}})"
        )
        cmd = [sys.executable, "-c", translate_script]

        env = {
            **os.environ,
            "OPENAI_BASE_URL": "https://api.deepseek.com",
            "OPENAI_API_KEY": DEEPSEEK_API_KEY,
            "OPENAI_MODEL": "deepseek-chat",
            "HF_ENDPOINT": "https://hf-mirror.com",
        }

        yield f"data: {json.dumps({'stage': 'translating', 'message': '正在翻译论文（可能需要 2-5 分钟）...'}, ensure_ascii=False)}\n\n"
        await asyncio.sleep(0.1)

        try:
            process = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: subprocess.run(
                    cmd, env=env, capture_output=True, text=True, timeout=600
                )
            )

            if process.returncode != 0:
                error_msg = process.stderr[:500] if process.stderr else "未知错误"
                logger.error(f"pdf2zh 翻译失败: {error_msg}")
                yield f"data: {json.dumps({'stage': 'error', 'message': f'翻译失败: {error_msg}'}, ensure_ascii=False)}\n\n"
                return

            # 查找输出文件
            result = _find_translated_files(translated_dir, base_name)
            if not result:
                yield f"data: {json.dumps({'stage': 'error', 'message': '翻译完成但未找到输出文件'}, ensure_ascii=False)}\n\n"
                return

            yield f"data: {json.dumps({'stage': 'done', 'message': '翻译完成！', **result}, ensure_ascii=False)}\n\n"

        except subprocess.TimeoutExpired:
            yield f"data: {json.dumps({'stage': 'error', 'message': '翻译超时（超过10分钟），请稍后重试'}, ensure_ascii=False)}\n\n"
        except Exception as e:
            logger.exception("翻译异常")
            yield f"data: {json.dumps({'stage': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


def _find_translated_files(translated_dir: str, base_name: str) -> dict:
    """查找翻译后的 PDF 文件，返回 URL dict 或空 dict"""
    if not os.path.isdir(translated_dir):
        return {}

    result = {}
    for f in os.listdir(translated_dir):
        if f.endswith(".pdf"):
            url = f"/api/uploads/{base_name}_translated/{f}"
            fl = f.lower()
            if "dual" in fl or "bilingual" in fl:
                result["dual_pdf_url"] = url
            elif "mono" in fl or "zh" in fl or "translated" in fl:
                result["mono_pdf_url"] = url
            else:
                # 兜底：任何 PDF 都归为 mono
                if "mono_pdf_url" not in result:
                    result["mono_pdf_url"] = url
    return result



class CodeGenerateRequest(BaseModel):
    user_request: str = "生成论文核心算法的完整实现代码"
    target_framework: str = "Python (PyTorch)"


@app.post("/api/code/generate")
async def generate_code_stream(request: CodeGenerateRequest):
    """流式生成论文代码复现"""
    if not app_state.is_document_loaded:
        raise HTTPException(status_code=400, detail="请先上传并解析论文文档")

    coordinator = get_coordinator()

    async def generate():
        try:
            for chunk in coordinator.generate_code_stream(
                user_request=request.user_request,
                target_framework=request.target_framework,
            ):
                yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
                await asyncio.sleep(0.01)

            yield f"data: {json.dumps({'done': True}, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


@app.post("/api/clear")
def clear_chat():
    """清除聊天历史"""
    coordinator = get_coordinator()
    coordinator.clear_chat_history()

    # 同时清除持久化的对话记录
    if app_state.current_document_id:
        try:
            get_history_store().clear_chat_history(app_state.current_document_id)
        except Exception as e:
            logger.error(f"清除历史记录失败: {e}")

    return {"success": True}


@app.delete("/api/document")
def clear_document():
    """清除当前文档"""
    coordinator = get_coordinator()
    coordinator.clear_chat_history()
    if app_state.current_document_id:
        app_state.remove_document(app_state.current_document_id)
    else:
        app_state.clear()
    return {"success": True}


# ===================== 论文对比分析 =====================

class CompareRequest(BaseModel):
    doc_ids: List[str]

@app.post("/api/compare/stream")
async def compare_stream(request: CompareRequest):
    """流式对比分析多篇论文"""
    doc_ids = request.doc_ids
    if len(doc_ids) < 2:
        raise HTTPException(status_code=400, detail="至少选择 2 篇论文进行对比")
    if len(doc_ids) > 3:
        raise HTTPException(status_code=400, detail="最多支持 3 篇论文对比")

    papers = []
    for did in doc_ids:
        if did not in app_state.documents:
            raise HTTPException(status_code=404, detail=f"文档 {did} 不存在")
        doc_data = app_state.documents[did]
        summary = doc_data.get("summary", "")
        if not summary:
            info = doc_data.get("document_info", {})
            raise HTTPException(status_code=400, detail=f"文档「{info.get('filename', did)}」尚未分析")
        info = doc_data.get("document_info", {})
        papers.append({
            "title": info.get("title") or info.get("filename", f"论文{len(papers)+1}"),
            "summary": summary[:6000],
        })

    papers_text = ""
    for i, p in enumerate(papers):
        papers_text += f"\n{'='*50}\n论文 {i+1}: {p['title']}\n{'='*50}\n{p['summary']}\n"

    col3 = " 论文3 |" if len(papers) > 2 else ""
    sep3 = "-------|" if len(papers) > 2 else ""
    prompt = f"""你是资深学术研究员，请对以下 {len(papers)} 篇论文进行深度对比分析。
{papers_text}

请按以下结构输出（Markdown 格式）：

## 📊 对比概览表

| 维度 | 论文1 | 论文2 |{col3}
|------|-------|-------|{sep3}
| 研究目标 | | |
| 核心方法 | | |
| 数据集 | | |
| 主要结果 | | |
| 创新点 | | |

## 🔬 方法对比
详细对比各论文的技术路线差异。

## 📈 实验与结果对比
对比实验设计、评估指标和结果。

## 💡 创新点与贡献对比
分析各论文的独特贡献。

## ⚠️ 局限性对比
各论文的不足和局限。

## 🎯 总结与建议
综合评价异同，给出阅读建议。

请使用中文，对专业术语附英文原文。"""

    coordinator = get_coordinator()

    async def generate():
        try:
            full_text = ""
            for chunk in coordinator.llm_service.stream_chat(prompt):
                full_text += chunk
                yield f"data: {json.dumps({'stage': 'analyzing', 'chunk': chunk}, ensure_ascii=False)}\n\n"
                await asyncio.sleep(0.01)
            yield f"data: {json.dumps({'stage': 'done', 'analysis': full_text}, ensure_ascii=False)}\n\n"
        except Exception as e:
            logger.exception("对比分析失败")
            yield f"data: {json.dumps({'stage': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "Connection": "keep-alive"})


@app.get("/api/documents")
async def list_documents():
    """获取所有已加载的文档列表"""
    docs = []
    for doc_id, doc_data in app_state.documents.items():
        info = doc_data.get("document_info", {})
        docs.append({
            "document_id": doc_id,
            "filename": info.get("filename", ""),
            "title": info.get("title", ""),
            "file_url": info.get("file_url", ""),
            "page_count": info.get("page_count", 0),
            "is_active": doc_id == app_state.current_document_id,
            "has_summary": bool(doc_data.get("summary")),
        })
    return {"documents": docs, "active_id": app_state.current_document_id}


class SwitchDocumentRequest(BaseModel):
    document_id: str


@app.post("/api/documents/switch")
def switch_document(request: SwitchDocumentRequest):
    """切换活跃文档"""
    doc_id = request.document_id

    if doc_id not in app_state.documents:
        raise HTTPException(status_code=404, detail="文档不存在")

    if doc_id == app_state.current_document_id:
        return {
            "success": True, 
            "message": "已经是当前文档",
            "document_info": app_state.document_info
        }

    success = app_state.switch_to(doc_id)
    if not success:
        raise HTTPException(status_code=500, detail="切换失败")

    # 恢复 coordinator 上下文（set_document_context 内部已含 load_collection）
    coordinator = get_coordinator()
    doc_info = app_state.document_info
    coordinator.qa_agent.set_document_context(
        doc_id=doc_id,
        paper_title=doc_info.get("title", ""),
        paper_summary=app_state.current_summary,
    )
    # 恢复 coordinator 的 current_state
    if coordinator.current_state:
        coordinator.current_state["summary"] = app_state.current_summary

    return {
        "success": True,
        "document_info": doc_info,
        "message": f"已切换到: {doc_info.get('filename', '')}"
    }


@app.delete("/api/documents/{doc_id}")
def remove_document(doc_id: str):
    """移除指定文档"""
    if doc_id not in app_state.documents:
        raise HTTPException(status_code=404, detail="文档不存在")

    was_active = doc_id == app_state.current_document_id
    app_state.remove_document(doc_id)

    # 如果删除的是当前文档，恢复新活跃文档的上下文
    if was_active and app_state.is_document_loaded:
        coordinator = get_coordinator()
        new_id = app_state.current_document_id
        doc_info = app_state.document_info
        coordinator.qa_agent.set_document_context(
            doc_id=new_id,
            paper_title=doc_info.get("title", ""),
            paper_summary=app_state.current_summary,
        )

    return {"success": True}


@app.get("/api/history", response_model=HistoryListResponse)
async def get_analysis_history():
    """获取分析历史记录列表"""
    try:
        history_list = get_history_store().get_analysis_history_list()
        history_items = [
            HistoryItem(
                id=h["id"],
                filename=h["filename"],
                title=h.get("title", ""),
                file_type=h.get("file_type", ""),
                page_count=h.get("page_count", 0),
                word_count=h.get("word_count", 0),
                processing_time=h.get("processing_time", 0),
                analyzed_at=h.get("analyzed_at", ""),
            )
            for h in history_list
        ]
        return HistoryListResponse(
            history=history_items, current_id=app_state.current_history_id
        )
    except Exception as e:
        logger.error("获取历史记录失败: %s", e)
        return HistoryListResponse(history=[], current_id=None)


@app.get("/api/history/{history_id}")
async def get_history_detail(history_id: str):
    """获取指定历史记录详情"""
    try:
        item = get_history_store().get_analysis_history_detail(history_id)
        if not item:
            raise HTTPException(status_code=404, detail="历史记录不存在")
        return item
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取详情失败: {e}")
        raise HTTPException(status_code=500, detail="服务错误")


@app.post("/api/history/{history_id}/load")
def load_history_item(history_id: str):
    """加载历史记录到当前状态"""
    coordinator = get_coordinator()

    try:
        history_store = get_history_store()
        item = history_store.get_analysis_history_detail(history_id)
        if not item:
            raise HTTPException(status_code=404, detail="历史记录不存在")

        # 更新当前状态
        app_state.current_summary = item.get("summary", "")
        app_state.current_structure = item.get("structure", "")
        app_state.is_document_loaded = True
        app_state.current_history_id = history_id
        app_state.current_document_id = item.get("document_id", "")
        app_state.document_info = {
            "filename": item["filename"],
            "title": item.get("title", ""),
            "file_type": item.get("file_type", ""),
            "page_count": item.get("page_count", 0),
            "word_count": item.get("word_count", 0),
            "processing_time": item.get("processing_time", 0),
            "document_id": item.get("document_id", ""),
        }

        # 尝试加载向量存储（set_document_context 内部已含 load_collection）
        if app_state.current_document_id:
            coordinator.qa_agent.set_document_context(
                doc_id=app_state.current_document_id,
                paper_title=item.get("title", ""),
                paper_summary=app_state.current_summary[:500],
            )

        # 获取该文档的对话历史
        chat_history = history_store.get_chat_history(app_state.current_document_id)

        return {
            "success": True,
            "document_info": app_state.document_info,
            "structure": app_state.current_structure,
            "summary": app_state.current_summary,
            "chat_history": chat_history,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"加载历史记录失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/history/chat")
async def get_current_history_chat():
    """获取当前活跃文档的对话历史"""
    if not app_state.current_document_id:
        return {"chat_history": []}
    try:
        chat_history = get_history_store().get_chat_history(app_state.current_document_id)
        return {"chat_history": chat_history}
    except Exception as e:
        logger.error(f"获取当前对话历史失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/document/profile", response_model=ProfileSummaryResponse)
def get_document_profile_summary():
    """获取当前文档的结构化 profile 摘要"""
    profile = _load_profile_for_document(app_state.current_document_id or "")
    if not profile:
        raise HTTPException(status_code=404, detail="当前文档暂无结构化 profile")
    return ProfileSummaryResponse(**_build_profile_summary(profile))


@app.get("/api/history/{history_id}/profile-summary", response_model=ProfileSummaryResponse)
def get_history_profile_summary(history_id: str):
    """获取指定历史记录对应论文的结构化 profile 摘要"""
    history_store = get_history_store()
    item = history_store.get_analysis_history_detail(history_id)
    if not item:
        raise HTTPException(status_code=404, detail="历史记录不存在")
    profile = _load_profile_for_document(item.get("document_id", ""))
    if not profile:
        raise HTTPException(status_code=404, detail="该历史记录暂无结构化 profile")
    return ProfileSummaryResponse(**_build_profile_summary(profile))


@app.get("/api/document/profile/detail", response_model=ProfileDetailResponse)
def get_document_profile_detail():
    """获取当前文档的完整结构化 profile"""
    profile = _load_profile_for_document(app_state.current_document_id or "")
    if not profile:
        raise HTTPException(status_code=404, detail="当前文档暂无结构化 profile")
    return ProfileDetailResponse(**_build_profile_detail(profile))


@app.get("/api/history/{history_id}/profile/detail", response_model=ProfileDetailResponse)
def get_history_profile_detail(history_id: str):
    """获取指定历史记录对应论文的完整结构化 profile"""
    history_store = get_history_store()
    item = history_store.get_analysis_history_detail(history_id)
    if not item:
        raise HTTPException(status_code=404, detail="历史记录不存在")
    profile = _load_profile_for_document(item.get("document_id", ""))
    if not profile:
        raise HTTPException(status_code=404, detail="该历史记录暂无结构化 profile")
    return ProfileDetailResponse(**_build_profile_detail(profile))


@app.delete("/api/history/{history_id}")
def delete_history_item(history_id: str):
    """删除指定历史记录"""
    coordinator = get_coordinator()

    try:
        history_store = get_history_store()
        item = history_store.get_analysis_history_detail(history_id)
        if not item:
            raise HTTPException(status_code=404, detail="历史记录不存在")

        # 删除历史记录
        success = history_store.delete_analysis_history(history_id)
        if not success:
            raise HTTPException(status_code=500, detail="删除失败")

        # 尝试删除向量集合
        if item.get("document_id"):
            try:
                coordinator.vector_store.delete_collection(item["document_id"])
            except Exception as e:
                logger.warning(f"删除向量集合失败: {e}")

        # 如果删除的是当前记录，清除当前状态
        if app_state.current_history_id == history_id:
            app_state.clear()

        return {"success": True}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"删除失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/history/{history_id}/chat")
async def get_history_chat(history_id: str):
    """获取指定历史记录的对话历史"""
    try:
        history_store = get_history_store()
        item = history_store.get_analysis_history_detail(history_id)
        if not item:
            raise HTTPException(status_code=404, detail="历史记录不存在")

        document_id = item.get("document_id", "")
        chat_history = history_store.get_chat_history(document_id)

        return {"chat_history": chat_history}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取对话历史失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== 课题组讨论 ====================

_lab_session_cache: Optional[LabSession] = None


@app.get("/api/lab/proposal")
def get_lab_proposal():
    """获取最近一次课题组讨论的研究提案"""
    global _lab_session_cache
    if _lab_session_cache and _lab_session_cache.proposal:
        return {
            "success": True,
            "proposal": _lab_session_cache.proposal,
            "discussion": _lab_session_cache.shared_memory.get_discussion_text(max_chars=10000),
        }
    return {"success": False, "message": "暂无课题组讨论结果"}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8001)