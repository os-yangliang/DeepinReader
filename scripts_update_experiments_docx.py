from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
SRC=r"d:\PaperReader\PaperReader\docs\CE-RAG-conference-draft.docx"
OUT=r"d:\PaperReader\PaperReader\docs\CE-RAG-conference-draft-submission.docx"

def cut(doc,title):
    b=doc._body._element; xs=list(b); k=None
    for i,e in enumerate(xs):
        s=''.join(n.text for n in e.iter() if n.tag.endswith('}t') and n.text).strip()
        if s==title: k=i; break
    if k is not None:
        for e in xs[k:]: b.remove(e)

def p(d,s): d.add_paragraph(s)
def h(d,s,l): d.add_heading(s,l)
def tb(d,cap,heads,rows):
    q=d.add_paragraph(cap); q.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in q.runs: r.italic=True
    t=d.add_table(rows=1,cols=len(heads)); t.style='Table Grid'
    for i,x in enumerate(heads): t.rows[0].cells[i].text=x
    for row in rows:
        cs=t.add_row().cells
        for i,x in enumerate(row): cs[i].text=str(x)
    d.add_paragraph()

d=Document(SRC); cut(d,'4 Experiments')
h(d,'4 Experiments',1)
p(d,'This section evaluates CE-RAG for faithful scientific paper question answering. The experiments address four research questions: whether CE-RAG improves answer quality and faithfulness over conventional RAG baselines; whether reasoning-chain retrieval improves evidence coverage and interpretability; whether evidence sufficiency estimation and verification reduce unsupported answers; and how CE-RAG performs across different question types. Numerical values are left as TBD placeholders and will be completed after the final experimental run.')
h(d,'4.1 Dataset Construction',2)
p(d,'We construct a pilot scientific paper QA dataset for evaluating claim-evidence-aware paper question answering. The dataset contains 53 scientific papers from the computer science domain and 318 question-answer pairs. For each paper, six question templates are generated to cover method understanding, evidence analysis, result interpretation, contribution summarization, limitation analysis, and unanswerable questions. Unanswerable questions are included to evaluate whether a system can avoid over-generalized conclusions when the paper does not provide sufficient evidence.')
p(d,'The dataset is built through an LLM-assisted annotation pipeline followed by manual review. Each paper is parsed into text chunks, and candidate contexts are selected according to question type and keyword matching. The annotation model is instructed to generate a concise reference answer and extract supporting evidence only from the provided paper context. Evidence snippets must be grounded in the original paper text, and questions without sufficient support are marked as unanswerable.')
tb(d,'Table 1. Statistics of the constructed QA dataset.',['Item','Value'],[['Papers','53'],['QA pairs','318'],['Questions per paper','6'],['Domain','Computer science'],['Question types','Method, Evidence, Result, Contribution, Limitation, Unanswerable'],['Annotation','LLM-assisted annotation with manual review'],['Average paper length','TBD'],['Average answer length','TBD']])
tb(d,'Table 2. Question types used in the dataset.',['Type','Purpose','Example'],[['Method','Identify the core method','What is the core method proposed in this paper?'],['Evidence','Ask how claims are supported','How do the authors demonstrate effectiveness?'],['Result','Summarize experimental findings','What are the key experimental results?'],['Contribution','Summarize contributions','What are the main contributions?'],['Limitation','Identify limitations','What are the limitations or future directions?'],['Unanswerable','Test unsupported over-generalization','Does the method outperform all existing methods on all tasks?']])
h(d,'4.2 Baselines',2)
p(d,'We compare CE-RAG with two retrieval-augmented baselines. All methods use the same paper collection and QA samples. The comparison isolates the effect of scholarly object graphs, reasoning-chain retrieval, and evidence sufficiency estimation.')
tb(d,'Table 3. Compared methods.',['Method','Description'],[['Chunk-based RAG','Retrieves top-ranked text chunks using semantic similarity and generates answers from retrieved chunks.'],['Hybrid RAG','Combines dense retrieval with lightweight keyword or structural matching, without claim-evidence reasoning chains.'],['CE-RAG','Uses scholarly object graphs, claim-evidence reasoning chains, sufficiency estimation, and answer verification.']])
h(d,'4.3 Evaluation Metrics',2)
p(d,'Scientific paper QA should be evaluated by both answer quality and evidence reliability. We use correctness, faithfulness, evidence coverage, relevance, readability, hallucination rate, reasoning chain quality, and sufficiency judgment. Most metrics are scored on a 1-5 scale, while hallucination is treated as a binary indicator of unsupported or over-generalized statements. Final scores are averaged over evaluated samples.')
tb(d,'Table 4. Evaluation metrics.',['Metric','Scale','Description'],[['Correctness','1-5','Whether the answer correctly addresses the question.'],['Faithfulness','1-5','Whether the answer is supported by paper evidence.'],['Evidence Coverage','1-5','Whether key evidence is covered.'],['Relevance','1-5','Whether the answer is relevant.'],['Readability','1-5','Whether the answer is fluent and clear.'],['Hallucination Rate','0/1','Whether unsupported statements appear.'],['Reasoning Chain Quality','1-5','Whether retrieved chains provide useful support.'],['Sufficiency Judgment','1-5','Whether evidence sufficiency is judged correctly.']])
h(d,'4.4 Implementation Details',2)
p(d,'All methods are implemented in the same experimental framework. Each paper is parsed and indexed before answering questions. For chunk-based RAG and hybrid RAG, relevant text chunks are retrieved and used as context for generation. For CE-RAG, the paper is additionally transformed into a scholarly object graph containing sections, claims, evidence, experiments, results, and limitations. Candidate reasoning chains are generated by bounded-hop traversal from question-dependent anchors and ranked by relevance, type matching, edge confidence, evidence strength, and completeness.')
tb(d,'Table 5. Implementation settings.',['Parameter','Value'],[['Top-k retrieved chunks','TBD'],['Maximum reasoning-chain hops','TBD'],['Top-k reasoning chains','TBD'],['LLM for generation','TBD'],['Embedding model','TBD'],['Vector store','TBD'],['Sufficiency threshold','TBD'],['Environment','TBD']])
h(d,'4.5 Main Results',2)
p(d,'Table 6 presents the main comparison results. The comparison focuses on whether CE-RAG achieves higher faithfulness and evidence coverage while reducing hallucination rate. Since CE-RAG retrieves explicit reasoning chains and estimates evidence sufficiency, it is expected to provide more reliable evidence-grounded answers than chunk-based retrieval methods.')
tb(d,'Table 6. Main experimental results.',['Method','Correctness','Faithfulness','Evidence Coverage','Hallucination Rate','Chain Quality'],[['Chunk-based RAG','TBD','TBD','TBD','TBD','N/A'],['Hybrid RAG','TBD','TBD','TBD','TBD','N/A'],['CE-RAG','TBD','TBD','TBD','TBD','TBD']])
h(d,'4.6 Analysis by Question Type',2)
p(d,'We analyze performance across question types because different questions require different reasoning behaviors. Method and contribution questions can often be answered from high-level sections, whereas evidence and result questions require locating supporting experiments and results. Limitation and unanswerable questions are more challenging because they require recognizing missing or insufficient evidence.')
rows=[]
for qt in ['Method','Evidence','Result','Contribution','Limitation','Unanswerable']:
    for m in ['Chunk-based RAG','Hybrid RAG','CE-RAG']: rows.append([qt,m,'TBD','TBD','TBD','TBD'])
tb(d,'Table 7. Performance by question type.',['Question Type','Method','Correctness','Faithfulness','Evidence Coverage','Hallucination Rate'],rows)
h(d,'4.7 Ablation Study',2)
p(d,'To understand the contribution of each component, we remove reasoning-chain retrieval, evidence sufficiency estimation, and answer verification respectively. This study examines whether structured retrieval and reliability mechanisms contribute to faithful answer generation.')
tb(d,'Table 8. Ablation study.',['Variant','Correctness','Faithfulness','Evidence Coverage','Hallucination Rate'],[['CE-RAG w/o Reasoning Chains','TBD','TBD','TBD','TBD'],['CE-RAG w/o Sufficiency Estimation','TBD','TBD','TBD','TBD'],['CE-RAG w/o Verification','TBD','TBD','TBD','TBD'],['Full CE-RAG','TBD','TBD','TBD','TBD']])
h(d,'4.8 Case Study',2)
p(d,'We include a representative case study to illustrate how CE-RAG supports answer generation with explicit reasoning chains. The final case will be selected from experimental outputs and will compare baseline answers with the CE-RAG answer, retrieved reasoning chain, sufficiency score, and verifier warnings.')
tb(d,'Table 9. Case study template.',['Item','Content'],[['Question','TBD'],['Gold Answer','TBD'],['Chunk-based RAG Answer','TBD'],['Hybrid RAG Answer','TBD'],['CE-RAG Answer','TBD'],['Retrieved Reasoning Chain','Claim: TBD -> Evidence: TBD -> Result: TBD'],['Sufficiency Score','TBD'],['Verifier Warnings','TBD']])
h(d,'5 Discussion',1); p(d,'TBD. This section will discuss the main findings, failure cases, limitations of object extraction, and future extensions such as learnable reasoning-chain reranking and larger-scale evaluation.')
h(d,'6 Conclusion',1); p(d,'TBD. This section will summarize CE-RAG and highlight the benefits of claim-evidence reasoning chains and evidence sufficiency estimation.')
h(d,'References',1); p(d,'TBD')
d.save(OUT); print(OUT)
